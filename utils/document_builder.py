from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import math2docx
import os
from utils.download_file import download_file
from utils.authorization import _get_bearer_token
from utils.img_dimensions import img_dimensions
import re
from utils.logger import get_logger


def _sanitize_latex_for_docx(latex: str) -> str:
    r"""Fix latex2mathml bugs before passing to math2docx.

    Several LaTeX accent/grouping commands always produce malformed OMML XML
    (groupChrPr/groupChr tag mismatch). Their safe replacements:

        \bar     -> \overline      (same bar accent, wider variant works)
        \vec     -> \boldsymbol    (bold = standard vector notation fallback)
        \check   -> \widecheck     (same check accent, wider variant works)
        \widetilde -> \tilde       (same tilde, narrower variant works)
        \overrightarrow -> \boldsymbol
        \overleftarrow  -> \boldsymbol
        \overbrace{expr} -> expr   (strip grouping, keep content)
        \underbrace{expr} -> expr  (strip grouping, keep content)
    """
    latex = re.sub(r'\\bar\s*\{', r'\\overline{', latex)
    latex = re.sub(r'\\vec\s*\{', r'\\boldsymbol{', latex)
    latex = re.sub(r'\\check\s*\{', r'\\widecheck{', latex)
    latex = re.sub(r'\\widetilde\s*\{', r'\\tilde{', latex)
    latex = re.sub(r'\\overrightarrow\s*\{', r'\\boldsymbol{', latex)
    latex = re.sub(r'\\overleftarrow\s*\{', r'\\boldsymbol{', latex)
    # Strip \overbrace and \underbrace but preserve their content
    latex = re.sub(r'\\overbrace\s*\{', r'{', latex)
    latex = re.sub(r'\\underbrace\s*\{', r'{', latex)
    return latex

logger = get_logger(__name__)


def parse_markdown_text(text):
    """
    Parse a text string with simple Markdown emphasis into styled segments.

    Supports only `**bold**` and `*italic*` inline patterns.

    Args:
        text (str): Raw text that may contain Markdown emphasis.

    Returns:
        list[dict]: A sequence of segments with `text`, `bold`, and `italic`
        flags, preserving input order.
    """
    # Parse **bold** and *italic* in text
    # Split on ** and * patterns
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)
    segments = []
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            segments.append({'text': part[2:-2], 'bold': True, 'italic': False})
        elif part.startswith('*') and part.endswith('*'):
            segments.append({'text': part[1:-1], 'bold': False, 'italic': True})
        else:
            segments.append({'text': part, 'bold': False, 'italic': False})
    return segments

def _set_section_vertical_alignment(section, value: str) -> None:
    """Set a section's vertical text alignment via the `w:vAlign` element
    ('center' | 'top' | 'both' | 'bottom'). Word centers the page content natively, which is
    far more reliable than padding the cover with a guessed number of empty paragraphs."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:vAlign')):
        sectPr.remove(existing)
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), value)
    sectPr.append(v_align)


def build_docx_from_dict(doc_dict, buffer, request, URL):
    """
    Build a DOCX document from a normalized dictionary representation.

    The function generates a cover, optional page break/columns, and body
    elements (headers, paragraphs, lists, tables, images, and equations).
    Images are downloaded using the request authorization header; when an image
    cannot be loaded, a placeholder caption is inserted.

    Args:
        doc_dict (dict): Structured document definition with metadata and
            sections.
        buffer (BytesIO): Writable in-memory buffer where the DOCX is saved.
        request: Incoming request object used to extract authorization.
        URL (str): Base URL for file-download operations.

    Returns:
        BytesIO: The input buffer positioned at the beginning with DOCX content.

    Raises:
        ValueError: If a table element is missing required headers or rows.
    """
    logger.info("=> Starting document generation ...")

    metadata_data = doc_dict.get("metadata", {})
    sections_data = doc_dict.get("sections", [])
    # Sort sections by index_element
    sections_data.sort(key=lambda x: x.get("index_element", 0))
    font = doc_dict.get("font", "Times New Roman")
    # Document style drives the per-style formatting; columns are derived from it upstream.
    style_doc = str(doc_dict.get("style_doc", "report")).strip().lower()
    if style_doc not in ("ieee", "report"):
        style_doc = "report"
    is_ieee = style_doc == "ieee"
    columns_body = doc_dict.get("columns_body", 1)
    columns_body = int(columns_body)
    if columns_body > 2:
        columns_body = 2  # Limitar a máximo 2 columnas
    elif columns_body < 1:
        columns_body = 1  # Mínimo 1 columna

    doc = Document()

    # Aplicar metadata como portada centrada horizontalmente
    # Cover: horizontally centered, with proper point sizes so the title is centered and
    # correctly sized (the old Inches(0.5) ≈ 36pt read oversized and threw off the look).
    meta = metadata_data
    if "title" in meta:
        title = doc.add_paragraph(meta["title"].replace("*", ""))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(24)  # Título: grande pero proporcionado
            title.runs[0].font.name = font  # Usar font global
    if "subtitle" in meta:
        subtitle = doc.add_paragraph(meta["subtitle"].replace("*", ""))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle.runs:
            subtitle.runs[0].italic = True
            subtitle.runs[0].font.size = Pt(14)  # Subtítulo
            subtitle.runs[0].font.name = font
    if "description" in meta:
        desc = doc.add_paragraph(meta["description"].replace("*", ""))
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if desc.runs:
            desc.runs[0].font.size = Pt(11)
            desc.runs[0].font.name = font
    if "author" in meta:
        author = doc.add_paragraph(f"Autor: {meta['author'].replace('*','')}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if author.runs:
            author.runs[0].font.size = Pt(11)
            author.runs[0].font.name = font
    if "month" in meta and "year" in meta:
        date = doc.add_paragraph(f"{meta['month'].replace("*", "")} {meta['year'].replace("*", "")}")
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if date.runs:
            date.runs[0].font.size = Pt(11)
            date.runs[0].font.name = font
    # Page break + columns. When the cover gets its own page (page_break), isolate it in its
    # own section and let Word center it vertically (vAlign=center) — this centers the cover
    # correctly regardless of the title/font sizes. The body then starts top-aligned on a new
    # page (and in two columns for the IEEE style).
    page_break_requested = str(metadata_data.get("page_break", False)).lower() == "true"
    has_body = len(sections_data) > 0

    def _apply_columns(section):
        if columns_body > 1:
            cols = OxmlElement('w:cols')
            cols.set(qn('w:num'), str(columns_body))
            section._sectPr.append(cols)

    if page_break_requested:
        _set_section_vertical_alignment(doc.sections[0], 'center')  # center the cover page
        if has_body:
            body_section = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
            _set_section_vertical_alignment(body_section, 'top')    # body reads top-down
            _apply_columns(body_section)
    elif columns_body > 1:
        # No page break, but two columns: a continuous section break flows the body in columns.
        _apply_columns(doc.add_section(start_type=WD_SECTION_START.CONTINUOUS))
    
    # Contadores para numeración automática de captions
    figure_counter = 1
    table_counter = 1
    equation_counter = 1
    
    # Variable para acumular párrafos continuos
    current_paragraph = None
    
    # Procesar elementos
    for item in sections_data:
        # Detect type using the discriminator first if available
        item_type = item.get("type", None)
        
        if item_type in ("ParagraphHeader", "header") or ("text" in item and "level" in item):  # ParagraphHeader
            current_paragraph = None  # Reset paragraph
            level = item.get("level", 2)
            heading = doc.add_heading(item["text"], level=level)
            # Heading alignment per style: report → left (informe look); IEEE → level-1
            # section heads centered (article convention), deeper levels left-aligned.
            if is_ieee and level <= 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if heading.runs:
                heading.runs[0].font.name = font
        
        elif item_type in ("ParagraphBody", "paragraph") or (item_type is None and "text" in item and "bold" not in item):  # ParagraphBody
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # IEEE article style indents the first line of body paragraphs (incl. the
                # opening one). Report keeps paragraphs flush-left for a clean informe look.
                if is_ieee:
                    current_paragraph.paragraph_format.first_line_indent = Inches(0.2)
            segments = parse_markdown_text(item.get("text", ""))
            for seg in segments:
                run = current_paragraph.add_run(seg['text'])
                run.bold = seg['bold']
                run.italic = seg['italic']
                run.font.size = Inches(12 / 72)
                run.font.name = font

        elif item_type in ("ParagraphListItem", "list") or "items" in item:  # ParagraphListItem
            current_paragraph = None  # Reset paragraph
            list_style = item.get("list_style") or item.get("style") or "List Bullet"
            if list_style == "bullet":
                list_style = "List Bullet"
            elif list_style == "numbered":
                list_style = "List Number"
            for it in item["items"]:
                p = doc.add_paragraph(style='List Bullet' if list_style == "List Bullet" else 'List Number')
                segments = parse_markdown_text(it)
                for seg in segments:
                    run = p.add_run(seg['text'])
                    run.bold = seg['bold']
                    run.italic = seg['italic']
                    run.font.size = Inches(12 / 72)
                    run.font.name = font
        
        elif item_type in ("Table", "table") or "table_headers" in item or "headers" in item:  # Table
            current_paragraph = None  # Reset paragraph
            table_headers = item.get("table_headers") or item.get("headers")
            table_rows = item.get("table_rows") or item.get("rows")
            if not table_headers or not table_rows:
                raise ValueError("Table must have table_headers and table_rows.")
            caption_text = item.get("caption", f"Table {table_counter}: ")
            if not item.get("caption"):
                table_counter += 1
            p = doc.add_paragraph(caption_text, style='Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = 'Light List Accent 1'
            hdr_cells = table.rows[0].cells
            for i, hdr in enumerate(table_headers):
                hdr_cells[i].text = hdr
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.name = font
            for row_data in table_rows:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data
                    for run in row_cells[i].paragraphs[0].runs:
                        run.font.name = font
        
        elif item_type in ("Image", "image") or "id" in item:  # Image
            current_paragraph = None  # Reset paragraph
            try:
                bearer_token = _get_bearer_token(request)
                image_file = download_file(URL, bearer_token, item["id"])
                if isinstance(image_file, dict) and "error" in image_file:
                    raise ValueError(f"Error downloading image with ID {item['id']}: {image_file['error']['message']}")
                if not image_file or (hasattr(image_file, 'getbuffer') and len(image_file.getbuffer()) == 0):
                    raise ValueError(f"Downloaded image with ID {item['id']} is empty or invalid.")
                img_width, img_height = img_dimensions(image_file, body_columns=columns_body)
                doc.add_picture(image_file, width=Inches(img_width), height=Inches(img_height))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Failed to load image {item['id']}: {e}. Adding placeholder.")
                doc.add_paragraph(f"[Image Placeholder: {item.get('caption', 'No caption')}]", style='Caption')
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_text = item.get("caption", f"Figure {figure_counter}: ")
            if not item.get("caption"):
                figure_counter += 1
            p = doc.add_paragraph(caption_text, style='Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif item_type in ("Equation", "equation") or "latex" in item:  # Equation
            current_paragraph = None  # Reset paragraph
            p = doc.add_paragraph()
            math2docx.add_math(p, _sanitize_latex_for_docx(item["latex"]))
            if item.get("caption"):
                caption_text = item.get("caption", f"Equation {equation_counter}: ")
                if not item.get("caption"):
                    equation_counter += 1
                p_cap = doc.add_paragraph(caption_text, style='Caption')
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif item_type == "page_break":
            current_paragraph = None
            doc.add_page_break()

        else:
            current_paragraph = None  # Reset for any other item

    logger.info("=> Document generation completed!")
    doc.save(buffer)
    buffer.seek(0)
    return buffer