# Native libraries
from json import dumps
from os import getenv, makedirs
from os.path import exists
from shutil import rmtree
from typing import Annotated, Literal, List, Tuple
from enum import Enum
from uuid import uuid4
from pathlib import Path
from io import BytesIO

# Third-party libraries
from pydantic import Field
from requests import post, get
from mcp.server.fastmcp import FastMCP
from docx import Document

# from dotenv import load_dotenv
# load_dotenv()

URL = getenv('OWUI_URL')
TOKEN = getenv('JWT_SECRET')
PORT = int(getenv('PORT'))

# helpers uploading files
def upload_file(url: str, token: str, file_path: str, filename:str, file_type:str) -> dict:
    """ 
    Upload a file to the specified URL with the provided token.
    Args:
        url (str): The URL to which the file will be uploaded.
        token (str): The authorization token for the request.
        file_path (str): The path to the file to be uploaded.
        filename (str): The desired filename for the uploaded file.
        file_type (str): The file extension/type (e.g., 'pptx', 'xlsx', 'docx', 'md').
    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the uploaded file.
              Format: "[Download {filename}.{file_type}](/api/v1/files/{id}/content)"
              On error: {"error": {"message": "error description"}}
    """
    # Ensure the URL ends with '/api/v1/files/'
    url = f'{url}/api/v1/files/'

    # Prepare headers and files for the request
    headers = {
        'Authorization': f'Bearer {token}',
        'Accept': 'application/json'
    }

    # Open the file and send the POST request
    with open(file_path, 'rb') as f:
        files = {'file': f}
        response = post(url, headers=headers, files=files)


    if response.status_code != 200:
       return dumps({"error":{"message": f'Error uploading file: {response.status_code}'}})
    else:
        return dumps(
            {
            "file_path_download": f"[Download {filename}.{file_type}](/api/v1/files/{response.json()['id']}/content)"
            },
            indent=4,
            ensure_ascii=False
        )

# helpers downloading files in memory
def download_file(url: str, token: str, file_id: str) -> BytesIO:
    """
    Download a file from the specified URL with the provided token and file ID.
    Args:
        url (str): The base URL from which the file will be downloaded.
        token (str): The authorization token for the request.
        file_id (str): The ID of the file to be downloaded.
    """
    # Ensure the URL ends with '/api/v1/files/'
    url = f'{url}/api/v1/files/{file_id}/content'

    # Prepare headers and files for the request
    headers = {
        'Authorization': f'Bearer {token}',
        'Accept': 'application/json'
    }
    # Send the GET request
    response = get(url, headers=headers)

    if response.status_code != 200:
       return {"error":{"message": f'Error downloading the file: {response.status_code}'}}
    else:
        return BytesIO(response._content)
    

mcp = FastMCP(
    name="GenFilesMCP",
    instructions=
        "Generates PowerPoint, Excel, Word or Markdown files from user requests. Each tool returns a markdown hyperlink for downloading the generated file. Use the specific tools for each file type: generate_powerpoint, generate_excel, generate_word, or generate_markdown. For reviewing existing files, use full_context_docx to analyze structure and review_docx to add comments.",
    port=PORT,
    host="0.0.0.0"
)


@mcp.tool(
    name="generate_powerpoint",
    title="Generate PowerPoint presentation",
    description="""Generate a PowerPoint presentation using a Python script. Returns a markdown hyperlink for downloading the generated file.

Template structure:
```python
def power_point():
    # Allowed packages
    import numpy as np
    import os
    from pptx import Presentation

    # Import here other pptx packages you need, but do not import other packages that are not allowed.

    # Path to save the PowerPoint file, previously defined in the server.py file
    PPTX_PATH = pptx_path # Do not modify this line, it is defined in the server.py file

    # Initialize a new Presentation instance
    prs = Presentation() # slides ratio has to be 16:9 not 4:3

    # Generate here the necessary transformations for generating the PowerPoint presentation according to the user's request. Use titles, subtitles, diagrams, tables, colors, clear fonts, and other elements to make the presentation visually appealing and easy to understand.

    # Save the presentation
    prs.save(PPTX_PATH) # Do not modify this line, it is defined in the server.py file

    # Check if the file was created successfully
    if not os.path.exists(PPTX_PATH):
        raise ValueError(f"Failed to create the PowerPoint file at {PPTX_PATH}, try again")
    else:
        return f"PowerPoint file created successfully!"

# Invoke the function to generate the PowerPoint presentation
power_point()
```

Provide a complete Python script following this template to generate your PowerPoint presentation."""
)
def generate_powerpoint(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the PowerPoint presentation using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated PowerPoint file without the extension.")
    ]
) -> dict:
    """
    Generate a PowerPoint file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated PowerPoint file.
              Format: "[Download {filename}.pptx](/api/v1/files/{id}/content)"
    """
    # user folder
    if not exists('/app/temp'):
        makedirs('/app/temp')
    try:
        # Generate a unique filename for the PowerPoint file
        file_path = f'/app/temp/{file_name}_{uuid4()}.pptx'
        context = {"pptx_path": file_path}
        exec(python_script, context )

        # Upload the generated PowerPoint file
        response = upload_file(
            url=URL, 
            token=TOKEN, 
            file_path=file_path,
            filename=file_name,
            file_type="pptx"
        )
        # Response format: {"file_path_download": "[Download presentation.pptx](/api/v1/files/123/content)"}

        # remove the temporary file after upload
        rmtree('/app/temp', ignore_errors=True)
        
        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name="generate_excel",
    title="Generate Excel workbook",
    description="""Generate an Excel workbook using a Python script. Returns a markdown hyperlink for downloading the generated file.

Template structure:
```python
# Allowed packages
import numpy as np
import os
from openpyxl import Workbook

# Import here other xlsx packages you need, but do not import other packages that are not allowed.

# Path to save excel file, previously defined in the server.py file
XLSX_PATH = xlsx_path # Do not modify this line, it is defined in the server.py file

def excel():
    # Initialize a new Workbook instance
    wb = Workbook()

    # Apply the required data transformations to build the Excel workbook based on the user's request.
    # Create the necessary worksheets, populate tables, add charts, and format cells for clarity and visual appeal.

    # Save the Excel workbook
    wb.save(XLSX_PATH) # Do not modify this line, it is defined in the server.py file

    # Check if the file was created successfully
    if not os.path.exists(XLSX_PATH):
        raise ValueError(f"Failed to create the excel file.")
    else:
        return f"Excel file created successfully!"

# Invoke the function to generate the Excel file
excel()
```

Provide a complete Python script following this template to generate your Excel workbook."""
)
def generate_excel(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the Excel workbook using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated Excel file without the extension.")
    ]
) -> dict:
    """
    Generate an Excel file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Excel file.
              Format: "[Download {filename}.xlsx](/api/v1/files/{id}/content)"
    """
    # user folder
    if not exists('/app/temp'):
        makedirs('/app/temp')
    try:
        # Generate a unique filename for the Excel file
        file_path = f'/app/temp/{file_name}_{uuid4()}.xlsx'
        context = {"xlsx_path": file_path}
        exec(python_script, context )

        # Upload the generated Excel file
        response = upload_file(
            url=URL, 
            token=TOKEN, 
            file_path=file_path,
            filename=file_name,
            file_type="xlsx"
        )
        # Response format: {"file_path_download": "[Download workbook.xlsx](/api/v1/files/123/content)"}

        # remove the temporary file after upload
        rmtree('/app/temp', ignore_errors=True)
        
        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name="generate_word",
    title="Generate Word document",
    description="""Generate a Word document using a Python script. Returns a markdown hyperlink for downloading the generated file.

Template structure:
```python
def word():
    # Allowed packages
    import numpy as np
    import os
    from docx import Document

    # Import here other docx packages you need, but do not import other packages that are not allowed.

    # Path to save the docx file, previously defined in the server.py file
    DOCX_PATH = docx_path # Do not modify this line, it is defined in the server.py file

    # Initialize a new Document instance
    doc = Document()

    # Generate here the necessary transformations for generating the word document to the user's request. 

    # Save the presentation
    doc.save(DOCX_PATH) # Do not modify this line, it is defined in the server.py file

    # Check if the file was created successfully
    if not os.path.exists(DOCX_PATH):
        raise ValueError(f"Failed to create the word file.")
    else:
        return f"Word file created successfully!"

# Invoke the function to generate the word document
word()
```

Provide a complete Python script following this template to generate your Word document."""
)
def generate_word(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the Word document using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated Word file without the extension.")
    ]
) -> dict:
    """
    Generate a Word file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Word file.
              Format: "[Download {filename}.docx](/api/v1/files/{id}/content)"
    """
    # user folder
    if not exists('/app/temp'):
        makedirs('/app/temp')
    try:
        # Generate a unique filename for the Word file
        file_path = f'/app/temp/{file_name}_{uuid4()}.docx'
        context = {"docx_path": file_path}
        exec(python_script, context )

        # Upload the generated Word file
        response = upload_file(
            url=URL, 
            token=TOKEN, 
            file_path=file_path,
            filename=file_name,
            file_type="docx"
        )
        # Response format: {"file_path_download": "[Download document.docx](/api/v1/files/123/content)"}

        # remove the temporary file after upload
        rmtree('/app/temp', ignore_errors=True)
        
        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name="generate_markdown",
    title="Generate Markdown document",
    description="""Generate a Markdown document using a Python script. Returns a markdown hyperlink for downloading the generated file.

Template structure:
```python
# Allowed packages
import numpy as np
import os
import pypandoc

# Import here other md packages you need, but do not import other packages that are not allowed.

# Path to save the Markdown file, previously defined in the server.py file
MD_PATH = md_path # Do not modify this line, it is defined in the server.py file

def markdown():
    # Build a Markdown document according to the user's request.
    markdown_content = "# Your Markdown Content Here"

    # Save the markdown
    pypandoc.convert_text(markdown_content, 'md', format='md', outputfile=MD_PATH, extra_args=['--standalone']) # Do not modify this line, it is defined in the server.py file

    # Check if the file was created successfully
    if not os.path.exists(MD_PATH):
        raise ValueError(f"Failed to create the Markdown file.")
    else:
        return f"Markdown file created successfully!"

# Invoke the function to generate the markdown document
markdown()
```

Provide a complete Python script following this template to generate your Markdown document."""
)
def generate_markdown(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the Markdown document using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated Markdown file without the extension.")
    ]
) -> dict:
    """
    Generate a Markdown file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Markdown file.
              Format: "[Download {filename}.md](/api/v1/files/{id}/content)"
    """
    # user folder
    if not exists('/app/temp'):
        makedirs('/app/temp')
    try:
        # Generate a unique filename for the Markdown file
        file_path = f'/app/temp/{file_name}_{uuid4()}.md'
        context = {"md_path": file_path}
        exec(python_script, context )

        # Upload the generated Markdown file
        response = upload_file(
            url=URL, 
            token=TOKEN, 
            file_path=file_path,
            filename=file_name,
            file_type="md"
        )
        # Response format: {"file_path_download": "[Download document.md](/api/v1/files/123/content)"}

        # remove the temporary file after upload
        rmtree('/app/temp', ignore_errors=True)
        
        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )
    
@mcp.tool(
    name="full_context_docx",
    title="Return the structure of a docx document",
    description="""Return the index, style and text of each element in a docx document. This includes paragraphs, headings, tables, images, and other components. The output is a JSON object that provides a detailed representation of the document's structure and content.
    The Agent will use this tool to understand the content and structure of the document before perform corrections (spelling, grammar, style suggestions, idea enhancements). Agent have to identify the index of each element to be able to add comments in the review_docx tool."""
)
def full_context_docx(
    file_id: Annotated[
        str, 
        Field(description="ID of the existing docx file to analyze (from a previous chat upload).")
    ],
    file_name: Annotated[
        str, 
        Field(description="The name of the original docx file")
    ]
) -> dict:
    """
    Return the structure of a docx document including index, style, and text of each element.
    Returns:
        dict: A JSON object with the structure of the document.
    """
    try:
        # Download in memory the docx file using the download_file helper
        docx_file = download_file(
            url=URL, 
            token=TOKEN, 
            file_id=file_id
        )

        if isinstance(docx_file, dict) and "error" in docx_file:
            return dumps(
                docx_file,
                indent=4,
                ensure_ascii=False
            )
        else:
            # Instantiate a Document object from the in-memory file
            doc = Document(docx_file)
            
            # Structure to return
            text_body = {
                "file_name": file_name,
                "file_id": file_id,
                "body": []
            }

            # list to store different parts of the document
            parts = []

            for idx, parts in enumerate(doc.paragraphs):
                # text of the paragraph
                text = parts.text.strip()

                if not text:
                    # skip empty paragraphs
                    continue  

                # style of the paragraph
                style = parts.style.name  
                text_body["body"].append({
                    "index": idx,
                    "style": style ,  # style.name
                    "text": text  # text
                })

            return dumps(
                text_body,
                indent=4,
                ensure_ascii=False
            )
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name="review_docx",
    title="Review and comment on docx document",
    description="""Review an existing docx document, perform corrections (spelling, grammar, style suggestions, idea enhancements), and add comments to cells. Returns a markdown hyperlink for downloading the reviewed file."""
)
def review_docx(
    file_id: Annotated[
        str, 
        Field(description="ID of the existing docx file to review (from a previous chat upload).")
    ],
    file_name: Annotated[
        str, 
        Field(description="The name of the original docx file")
    ],
    review_comments: Annotated[
        List[Tuple[int, str]], 
        Field(description="List of tuples where each tuple contains: (index: int - the index of the docx document element to comment on, comment: str - the review comment, idea enhancements, suggestions or correction text).")
    ]
) -> dict:
    """
    Review an existing docx document and add comments to specified elements.
    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the reviewed docx file.
              Format: "[Download {filename}.md](/api/v1/files/{id}/content)"
    """
    # user folder
    if not exists('/app/temp'):
        makedirs('/app/temp')
    try:
        
        # Download the existing docx file
        docx_file = download_file(URL, TOKEN, file_id)
        if isinstance(docx_file, dict) and "error" in docx_file:
            return dumps(docx_file, indent=4, ensure_ascii=False)

        # Load the document
        doc = Document(docx_file)

        # Add comments to specified paragraphs
        paragraphs = list(doc.paragraphs)  # Get list of paragraphs
        for index, comment_text in review_comments:
            if 0 <= index < len(paragraphs):
                para = paragraphs[index]
                if para.runs:  # Ensure there are runs to comment on
                    # Add comment to the first run of the paragraph
                    doc.add_comment(
                        runs=[para.runs[0]],
                        text=comment_text,
                        author="AI Reviewer",
                        initials="AI"
                    )

        # Save the reviewed file
        reviewed_path = f'/app/temp/{Path(file_name).stem}_reviewed_{uuid4()}.docx'
        doc.save(reviewed_path)

        # Upload the reviewed docx file
        response = upload_file(
            url=URL, 
            token=TOKEN, 
            file_path=reviewed_path,
            filename=f"{Path(file_name).stem}_reviewed",
            file_type="docx"
        )

        # Remove temp file
        rmtree('/app/temp', ignore_errors=True)

        return response
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )
    

# Initialize and run the server
if __name__ == "__main__":
    mcp.run(
        transport="streamable-http"
    )

